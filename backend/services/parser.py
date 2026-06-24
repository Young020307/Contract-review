import json
import re
from docx import Document
from docx.oxml.ns import qn


def _iter_body_paragraphs(doc):
    """Yield (text, runs, is_table_cell) for every paragraph in document body order.

    Includes both body <w:p> elements and <w:p> elements inside table cells.
    Each run is (text, is_underlined).
    """
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text, runs = _parse_paragraph_xml(child)
            if text:
                yield text, runs, False
        elif tag == 'tbl':
            for cell in child.iter(qn('w:tc')):
                for p in cell.findall(qn('w:p')):
                    text, runs = _parse_paragraph_xml(p)
                    if text:
                        yield text, runs, True


def _parse_paragraph_xml(p_el):
    """Extract text and runs from a <w:p> XML element.

    Uses iter() instead of findall() to include runs nested inside SDT
    (structured document tags), ins/del, and other wrapper elements, so
    the extracted text matches what mammoth.js produces in the frontend.
    """
    text_parts = []
    runs = []
    for r in p_el.iter(qn('w:r')):
        run_text = ''
        for t in r.findall(qn('w:t')):
            if t.text:
                run_text += t.text
        if run_text:
            rPr = r.find(qn('w:rPr'))
            is_under = rPr is not None and rPr.find(qn('w:u')) is not None
            runs.append((run_text, is_under))
            text_parts.append(run_text)
    full_text = ''.join(text_parts).strip()
    return full_text, runs


def _get_all_texts(doc):
    """Return flat list of paragraph texts in document body order (includes table cells)."""
    return [text for text, _, _ in _iter_body_paragraphs(doc)]


class DocxParser:
    @staticmethod
    def parse(file_path: str) -> list[dict]:
        doc = Document(file_path)
        paragraphs = []
        for i, (text, runs, is_table_cell) in enumerate(_iter_body_paragraphs(doc)):
            underline_ranges = []
            pos = 0
            for run_text, is_under in runs:
                run_len = len(run_text)
                if is_under:
                    underline_ranges.append([pos, pos + run_len])
                pos += run_len
            paragraphs.append({
                "index": i,
                "text": text,
                "underline_ranges": underline_ranges,
                "is_table_cell": is_table_cell
            })
        return paragraphs

    @staticmethod
    def extract_full_text(file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join(_get_all_texts(doc))

    @staticmethod
    def extract_fixed_text(file_path: str, annotations: list[dict]) -> str:
        doc = Document(file_path)
        all_texts = _get_all_texts(doc)
        ann_by_para = {}
        for a in annotations:
            ann_by_para.setdefault(a["paragraph_index"], []).append(a)

        texts = []
        for pi, text in enumerate(all_texts):
            if pi in ann_by_para:
                fillable_ranges = sorted(
                    [(a["start_char"], a["end_char"]) for a in ann_by_para[pi]
                     if a.get("zone_type") == "fillable"]
                )
                if not fillable_ranges:
                    texts.append(text)
                    continue
                fixed_parts = []
                pos = 0
                for start, end in fillable_ranges:
                    if pos < start:
                        fixed_parts.append(text[pos:start])
                    pos = max(pos, end)
                if pos < len(text):
                    fixed_parts.append(text[pos:])
                result = "".join(fixed_parts).strip()
                if result:
                    texts.append(result)
            else:
                texts.append(text)
        return "\n".join(texts)

    @staticmethod
    def _get_fixed_text_per_para(text: str, fillable_ranges: list[tuple[int, int]]) -> str:
        """Remove fillable zones from a single paragraph's text, returning the skeleton."""
        if not fillable_ranges:
            return text
        sorted_ranges = sorted(fillable_ranges, key=lambda r: r[0])
        fixed_parts = []
        pos = 0
        for start, end in sorted_ranges:
            if pos < start:
                fixed_parts.append(text[pos:start])
            pos = max(pos, end)
        if pos < len(text):
            fixed_parts.append(text[pos:])
        return "".join(fixed_parts)

    @staticmethod
    def align_paragraphs(tpl_paras: list[dict], doc_paras: list[dict],
                         annotations: list[dict]) -> dict:
        """Align template paragraphs to document paragraphs using fixed-text regex matching.

        For each template paragraph, extracts the fixed-text skeleton (non-fillable
        segments) and uses it as a search pattern to find the corresponding document
        paragraph. No similarity thresholds — matching is deterministic.

        Returns {"mapping": {tpl_idx: doc_idx|None}, "inserted": [doc_idx, ...]}
        """
        import re

        ann_by_para: dict[int, list[dict]] = {}
        for a in annotations:
            if a.get("zone_type") != "fillable":
                continue
            ann_by_para.setdefault(a["paragraph_index"], []).append(a)

        doc_texts = [p["text"] for p in doc_paras]
        used_docs: set[int] = set()
        mapping: dict[int, int | None] = {}
        deferred: list[int] = []  # tpl indices that need positional fallback

        # ── Pass 1: high-confidence matches ──
        for p in tpl_paras:
            pi = p["index"]
            text = p["text"]
            anns = ann_by_para.get(pi, [])
            found = None

            if not anns:
                # No fillable zones — exact, then substring, then high-similarity
                for j, dt in enumerate(doc_texts):
                    if j in used_docs:
                        continue
                    if dt == text:
                        found = j
                        break
                if found is None:
                    for j, dt in enumerate(doc_texts):
                        if j in used_docs:
                            continue
                        if text in dt or dt in text:
                            found = j
                            break
                if found is None:
                    from difflib import SequenceMatcher
                    best_j, best_r = None, 0.95
                    for j, dt in enumerate(doc_texts):
                        if j in used_docs:
                            continue
                        r = SequenceMatcher(None, text, dt).ratio()
                        if r > best_r:
                            best_j, best_r = j, r
                    found = best_j
            else:
                # Build fixed-text skeleton segments
                ranges = sorted(
                    [(a["start_char"], a["end_char"]) for a in anns],
                    key=lambda r: r[0])
                fixed_segments = []
                pos = 0
                for s, e in ranges:
                    s = max(0, min(s, len(text)))
                    e = max(s, min(e, len(text)))
                    if pos < s:
                        fixed_segments.append(text[pos:s])
                    pos = max(pos, e)
                if pos < len(text):
                    fixed_segments.append(text[pos:])

                if not fixed_segments:
                    # Entirely fillable — try text inclusion first
                    for j, dt in enumerate(doc_texts):
                        if j in used_docs:
                            continue
                        if text in dt or dt in text:
                            found = j
                            break
                    # If no text match, defer to pass 2
                else:
                    fixed_len = sum(len(s) for s in fixed_segments)
                    if fixed_len >= 2:
                        pattern = ".*?".join(re.escape(s) for s in fixed_segments)
                        for j, dt in enumerate(doc_texts):
                            if j in used_docs:
                                continue
                            if re.search(pattern, dt):
                                found = j
                                break
                    # If fixed text too short, defer to pass 2

            if found is not None:
                mapping[pi] = found
                used_docs.add(found)
            else:
                deferred.append(pi)

        # ── Pass 2: positional fallback for deferred paragraphs ──
        remaining_docs = [j for j in range(len(doc_texts)) if j not in used_docs]
        for pi in deferred:
            if remaining_docs:
                doc_j = remaining_docs.pop(0)
                mapping[pi] = doc_j
                used_docs.add(doc_j)
            else:
                mapping[pi] = None

        inserted = sorted(
            [j for j in range(len(doc_texts)) if j not in used_docs])

        return {"mapping": mapping, "inserted": inserted}

    @staticmethod
    def extract_fillable_values(template_file_path: str, doc_file_path: str, annotations: list[dict]) -> dict:
        """Extract fillable values using regex context anchors.

        Builds a regex per paragraph from template fixed-text segments surrounding
        the fillable zones, then extracts variable-length content from the filled
        document. Falls back to character-slice on regex mismatch.

        Returns dict keyed by '{pi}_{start_char}' with shape:
            {"value": str, "doc_start": int, "doc_end": int}
        """
        template_doc = Document(template_file_path)
        doc_doc = Document(doc_file_path)
        tpl_texts = _get_all_texts(template_doc)
        doc_texts = _get_all_texts(doc_doc)

        # Compute paragraph alignment
        tpl_paras = DocxParser.parse(template_file_path)
        doc_paras = DocxParser.parse(doc_file_path)
        alignment = DocxParser.align_paragraphs(tpl_paras, doc_paras, annotations)
        para_map = alignment["mapping"]

        values = {}

        ann_by_para: dict[int, list[dict]] = {}
        for ann in annotations:
            if ann.get("zone_type") != "fillable":
                continue
            pi = ann["paragraph_index"]
            ann_by_para.setdefault(pi, []).append(ann)

        for pi, anns in ann_by_para.items():
            if pi >= len(tpl_texts):
                continue
            doc_pi = para_map.get(pi)
            if doc_pi is None or doc_pi >= len(doc_texts):
                continue

            tpl_text = tpl_texts[pi]
            doc_text = doc_texts[doc_pi]
            if not tpl_text:
                continue

            anns_sorted = sorted(anns, key=lambda a: a.get("start_char", 0))
            tpl_len = len(tpl_text)

            fixed_parts = []
            pos = 0
            for ann in anns_sorted:
                start = max(0, min(ann.get("start_char", 0), tpl_len))
                end = max(start, min(ann.get("end_char", tpl_len), tpl_len))
                if pos < start:
                    fixed_parts.append(tpl_text[pos:start])
                else:
                    fixed_parts.append("")
                pos = max(pos, end)
            if pos < tpl_len:
                fixed_parts.append(tpl_text[pos:])
            else:
                fixed_parts.append("")

            escaped = [re.escape(p) for p in fixed_parts]
            pattern = "^" + "(.*?)".join(escaped) + "$"

            match = re.search(pattern, doc_text)
            if match:
                groups = match.groups()
                for i, ann in enumerate(anns_sorted):
                    key = f"{pi}_{ann.get('start_char', 0)}"
                    raw = groups[i] if i < len(groups) else ""
                    doc_start = match.start(i + 1) if i < len(groups) else 0
                    doc_end = match.end(i + 1) if i < len(groups) else 0
                    values[key] = {
                        "value": raw.strip().strip("_"),
                        "doc_start": doc_start,
                        "doc_end": doc_end
                    }
            else:
                for ann in anns_sorted:
                    tpl_start = max(0, min(ann.get("start_char", 0), len(doc_text)))
                    tpl_end = max(tpl_start, min(ann.get("end_char", len(doc_text)), len(doc_text)))
                    key = f"{pi}_{tpl_start}"
                    val = doc_text[tpl_start:tpl_end].strip().strip("_")
                    values[key] = {
                        "value": val,
                        "doc_start": tpl_start,
                        "doc_end": tpl_start + len(val)
                    }

        return values

    @staticmethod
    def detect_checkbox_status(template_file_path: str, doc_file_path: str,
                               annotations: list[dict]) -> dict:
        """Detect whether □ checkbox zones have been toggled to ☑ in the document.

        Only processes annotations whose rules include radio_group or checked=True.
        Returns dict keyed by '{pi}_{start_char}' with shape:
            {"checked": bool, "value": str, "doc_start": int, "doc_end": int}
        """
        CHECKED_CHARS = {'☑', '☒', '✓', '✔'}  # ☑ ☒ ✓ ✔

        template_doc = Document(template_file_path)
        doc_doc = Document(doc_file_path)
        tpl_texts = _get_all_texts(template_doc)
        doc_texts = _get_all_texts(doc_doc)

        # Compute paragraph alignment
        tpl_paras = DocxParser.parse(template_file_path)
        doc_paras = DocxParser.parse(doc_file_path)
        alignment = DocxParser.align_paragraphs(tpl_paras, doc_paras, annotations)
        para_map = alignment["mapping"]

        # Filter to checkbox-relevant annotations
        checkbox_anns = []
        for ann in annotations:
            if ann.get("zone_type") != "fillable":
                continue
            rules = ann.get("rules", {})
            if isinstance(rules, str):
                try:
                    rules = json.loads(rules)
                except (json.JSONDecodeError, TypeError):
                    rules = {}
            if not rules or not isinstance(rules, dict):
                continue
            if rules.get("radio_group"):
                checkbox_anns.append(ann)

        statuses = {}
        for ann in checkbox_anns:
            pi = ann["paragraph_index"]
            if pi >= len(tpl_texts):
                continue
            doc_pi = para_map.get(pi)
            if doc_pi is None or doc_pi >= len(doc_texts):
                continue

            start = max(0, ann.get("start_char", 0))
            end = max(start + 1, ann.get("end_char", start + 1))
            doc_text = doc_texts[doc_pi]

            if start < len(doc_text):
                doc_char = doc_text[start:min(end, len(doc_text))]
            else:
                doc_char = ""

            key = f"{pi}_{start}"
            statuses[key] = {
                "checked": bool(doc_char and all(c in CHECKED_CHARS for c in doc_char)),
                "value": doc_char,
                "doc_start": start,
                "doc_end": min(end, len(doc_text))
            }

        return statuses
